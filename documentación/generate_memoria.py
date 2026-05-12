"""
Script to generate the Outventura project memory document (Memoria de Proyecto).
Output: memoria_outventura.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Colors from the app theme ──────────────────────────────────────────────────
COLOR_PRIMARY       = RGBColor(0x58, 0x8C, 0x23)   # #588C23  green
COLOR_DARK_GREEN    = RGBColor(0x3B, 0x59, 0x3F)   # #3B593F
COLOR_LIGHT_GREEN   = RGBColor(0xAF, 0xD6, 0x8D)   # #AFD68D
COLOR_SECONDARY     = RGBColor(0xA6, 0x77, 0x4E)   # #A6774E  brown
COLOR_TERTIARY      = RGBColor(0x32, 0x47, 0x56)   # #324756  dark blue
COLOR_WHITE         = RGBColor(0xFF, 0xFF, 0xFF)
COLOR_DARK          = RGBColor(0x20, 0x18, 0x14)
COLOR_PLACEHOLDER   = RGBColor(0xCC, 0xCC, 0xCC)

doc = Document()

# ── Page margins ───────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(2.5)


# ── Helpers ────────────────────────────────────────────────────────────────────

def set_run_color(run, color):
    run.font.color.rgb = color

def add_heading(doc, text, level=1, color=None):
    """Add a heading with optional custom color."""
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if color:
        for run in p.runs:
            run.font.color.rgb = color
    return p

def add_body(doc, text, bold=False, italic=False, color=None, size=11):
    """Add a normal paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p

def add_image_placeholder(doc, description, height_cm=7):
    """Add a gray placeholder box indicating where an image should go."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Shading via paragraph border trick: we use a table with one cell
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    cell.width = Cm(14)
    # Set background color
    _set_cell_bg(cell, 'E8E8E8')
    # Set minimum height
    trPr = table.rows[0]._tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(int(height_cm * 567)))  # 1 cm ≈ 567 twips
    trHeight.set(qn('w:hRule'), 'atLeast')
    trPr.append(trHeight)

    p2 = cell.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p2.add_run(f"[ IMAGEN: {description} ]")
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    run.font.size = Pt(10)
    run.italic = True
    doc.add_paragraph()  # spacing after
    return table

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p

def add_numbered(doc, text):
    p = doc.add_paragraph(style='List Number')
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p

def _hex_from_rgb(color: RGBColor) -> str:
    """Return 6-char hex string from an RGBColor."""
    # RGBColor stores r, g, b as individual attributes via its __init__
    # but they're accessible as [0], [1], [2] or via str() = 'RRGGBB'
    s = str(color)  # e.g. '3B593F'
    return s.upper()

def _set_cell_bg(cell, hex_color: str):
    """Apply a background fill to a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_colored_table(doc, headers, rows, header_color=None):
    """Add a styled table."""
    hc = header_color or COLOR_DARK_GREEN
    hc_hex = _hex_from_rgb(hc)
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ''
        _set_cell_bg(cell, hc_hex)
        run = cell.paragraphs[0].add_run(h)
        run.font.color.rgb = COLOR_WHITE
        run.font.bold = True
        run.font.size = Pt(10)

    # Data rows
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        for ci, val in enumerate(row_data):
            row.cells[ci].text = str(val)
            for run in row.cells[ci].paragraphs[0].runs:
                run.font.size = Pt(10)
    doc.add_paragraph()
    return table


# ══════════════════════════════════════════════════════════════════════════════
#  PORTADA
# ══════════════════════════════════════════════════════════════════════════════

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Manuel Navalon Fornes")
run.font.size = Pt(13)
run.font.bold = True
run.font.color.rgb = COLOR_DARK_GREEN

doc.add_paragraph()
doc.add_paragraph()

add_image_placeholder(doc, "LOGO DE OUTVENTURA (icono montaña + texto)", height_cm=5)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("OUTVENTURA")
run.font.size = Pt(32)
run.font.bold = True
run.font.color.rgb = COLOR_PRIMARY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("L'aventura al teu abast")
run.font.size = Pt(14)
run.font.italic = True
run.font.color.rgb = COLOR_SECONDARY

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

for line in [
    ("CFGS DAW", 12, False),
    ("Mòdul Professional Projecte Intermodular", 12, False),
    ("Desenvolupament d'Aplicacions Multiplataforma", 12, False),
    ("IES L'ESTACIÓ", 12, True),
    ("Curs 2025-26", 11, False),
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(line[0])
    run.font.size = Pt(line[1])
    run.font.bold = line[2]
    run.font.color.rgb = COLOR_DARK_GREEN

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Alumne: Manuel Navalon Fornes")
run.font.size = Pt(12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Tutor: __________________________")
run.font.size = Pt(12)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
#  ÍNDEX
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "Índex", level=1, color=COLOR_PRIMARY)

index_entries = [
    "1.  Introducció",
    "    1.1  Idea de negoci",
    "    1.2  Objectius",
    "    1.3  Anàlisi del mercat i segmentació. Competència. DAFO",
    "    1.4  Alineació amb ODS",
    "2.  Anàlisi dels requeriments",
    "    2.1  Entitat Relació",
    "    2.2  Diagrama de classes i casos d'ús",
    "    2.3  Disseny",
    "        2.3.1  Mockups",
    "        2.3.2  Paleta de colors",
    "        2.3.3  Usabilitat",
    "    2.4  Tecnologies",
    "        2.4.1  Desplegament",
    "        2.4.2  Backend",
    "        2.4.3  Frontend",
    "        2.4.4  Disseny",
    "    2.5  Planificació",
    "        2.5.1  Diagrama de Gantt",
    "3.  Implementació",
    "    3.1  Desplegament",
    "    3.2  Backend",
    "    3.3  Frontend",
    "    3.4  Disseny",
    "4.  Millores futures",
    "5.  Conclusions",
    "    5.1  Personals",
    "    5.2  Tècniques",
    "6.  Referències bibliogràfiques",
]
for entry in index_entries:
    p = doc.add_paragraph(entry)
    p.paragraph_format.space_after = Pt(2)
    for run in p.runs:
        run.font.size = Pt(11)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
#  1. INTRODUCCIÓ
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "1. Introducció", level=1, color=COLOR_PRIMARY)

# 1.1
add_heading(doc, "1.1 Idea de negoci", level=2, color=COLOR_DARK_GREEN)
add_body(doc,
    "Outventura és una aplicació mòbil pensada per a empreses o associacions que "
    "organitzen activitats a l'aire lliure: excursions de muntanya, rutes aquàtiques, "
    "sortides de neu i acampades. La idea neix de la necessitat de digitalitzar la "
    "gestió d'aquestes activitats, que sovint es fa amb fulls de càlcul o amb paper."
)
add_body(doc,
    "L'aplicació permet als clients consultar el catàleg d'excursions disponibles, "
    "fer sol·licituds de participació i reservar material d'aventura (carpes, arnesos, "
    "cascos, etc.). Al mateix temps, els administradors i experts guies poden gestionar "
    "tot el flux: des de la creació de l'excursió fins al tancament de la reserva i el "
    "control de danys en el material retornat."
)
add_body(doc,
    "El model de negoci es basa en el cobro per excursió (preu base per participant) "
    "i per lloguer de material (preu diari per unitat). A més, s'aplica un càrrec "
    "addicional en cas de danys al material retornat."
)

# 1.2
add_heading(doc, "1.2 Objectius", level=2, color=COLOR_DARK_GREEN)
add_body(doc, "Els objectius principals del projecte són:")
add_bullet(doc, "Crear una app mòbil funcional per a iOS i Android amb Flutter.")
add_bullet(doc, "Desenvolupar un backend REST amb NestJS i una base de dades relacional PostgreSQL.")
add_bullet(doc, "Implementar un sistema de rols (SUPERADMIN, ADMIN, EXPERTO, USUARIO, INVITADO) amb autenticació JWT.")
add_bullet(doc, "Permetre la gestió completa del catàleg d'excursions i del material d'aventura.")
add_bullet(doc, "Facilitar el procés de sol·licitud i reserva per als clients.")
add_bullet(doc, "Tenir una UI moderna i coherent amb la identitat visual de l'app.")
add_bullet(doc, "Desplegar el backend amb Docker per facilitar la portabilitat.")

# 1.3
add_heading(doc, "1.3 Anàlisi del mercat i segmentació. Competència. DAFO", level=2, color=COLOR_DARK_GREEN)
add_body(doc,
    "El mercat de les activitats de turisme actiu a l'Estat espanyol ha crescut de "
    "manera notable en els últims anys. Empreses com Naturanea, Civitatis o Wikiloc "
    "cobreixen parcialment aquesta necessitat, però cap d'elles ofereix una solució "
    "completa per a la gestió interna d'una empresa d'activitats: control de material, "
    "assignació d'experts i seguiment d'estats."
)
add_body(doc, "Segmentació del mercat objectiu:")
add_bullet(doc, "Empreses d'activitats a l'aire lliure (guiatge, escalada, kayak...).")
add_bullet(doc, "Associacions esportives amb necessitat de gestionar material.")
add_bullet(doc, "Centres d'esplai o colònies que organitzen sortides periòdiques.")

add_body(doc, "")
add_body(doc, "Anàlisi DAFO:", bold=True)

dafo_headers = ["", "Positiu", "Negatiu"]
dafo_rows = [
    ["Intern", "Fortaleses:\n- Solució tot en un (gestió + client)\n- UI moderna i responsive\n- Arquitectura escalable", "Debilitats:\n- Backend en construcció\n- Equip d'un sol desenvolupador\n- Sense integració de pagament real"],
    ["Extern", "Oportunitats:\n- Mercat en creixement\n- Poca competència directa\n- Digitalització del sector", "Amenaces:\n- Apps genèriques com Google Forms\n- Cost d'adquisició de clients\n- Dependència de serveis cloud"],
]
add_colored_table(doc, dafo_headers, dafo_rows)

# 1.4
add_heading(doc, "1.4 Alineació amb objectius, valors, sostenibilitat i futur (ODS)", level=2, color=COLOR_DARK_GREEN)
add_body(doc,
    "Outventura s'alinea amb diversos Objectius de Desenvolupament Sostenible (ODS) "
    "de l'Agenda 2030 de les Nacions Unides:"
)
add_bullet(doc, "ODS 3 – Salut i benestar: promou l'activitat física a l'aire lliure.")
add_bullet(doc, "ODS 8 – Treball decent i creixement econòmic: digitalitza i professionalitza el sector del turisme actiu.")
add_bullet(doc, "ODS 11 – Ciutats i comunitats sostenibles: fomenta el turisme de proximitat i la connexió amb la natura.")
add_bullet(doc, "ODS 13 – Acció climàtica: l'app redueix el paper (gestió digital) i promou activitats de baix impacte ambiental.")

add_body(doc,
    "Pel que fa als valors de l'empresa, Outventura es basa en l'accessibilitat "
    "(activitats per a tots els nivells), la seguretat (control del material i dels guies) "
    "i el respecte pel medi natural."
)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
#  2. ANÀLISI DELS REQUERIMENTS
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "2. Anàlisi dels requeriments", level=1, color=COLOR_PRIMARY)

# 2.1
add_heading(doc, "2.1 Entitat Relació", level=2, color=COLOR_DARK_GREEN)
add_body(doc,
    "La base de dades de l'aplicació s'ha dissenyat amb PostgreSQL i gestionada "
    "mitjançant Prisma ORM. Les entitats principals del sistema són:"
)
add_bullet(doc, "Role: defineix els rols del sistema (SUPERADMIN, ADMIN, EXPERTO, USUARIO, INVITADO).")
add_bullet(doc, "User: usuaris del sistema, cadascun associat a un rol.")
add_bullet(doc, "Category: categories d'activitat (Aquàtic, Neu, Muntanya, Càmping).")
add_bullet(doc, "Activity: excursions o activitats disponibles al catàleg.")
add_bullet(doc, "Equipment: material d'aventura disponible per a lloguer.")
add_bullet(doc, "EquipmentStatus: estat del material (disponible, esgotat, manteniment, fora de servei).")

add_body(doc, "")
add_body(doc,
    "Les relacions principals són: un User té un Role, un Equipment té un EquipmentStatus, "
    "i tant Equipment com Activity poden tenir múltiples Categories (relació molts-a-molts)."
)

add_image_placeholder(doc, "DIAGRAMA ENTITAT-RELACIÓ: Role, User, Category, Activity, Equipment, EquipmentStatus amb totes les relacions i cardinalitats", height_cm=9)

# 2.2
add_heading(doc, "2.2 Diagrama de classes i casos d'ús", level=2, color=COLOR_DARK_GREEN)
add_body(doc,
    "El frontend utilitza una arquitectura neta (Clean Architecture) dividida en "
    "tres capes: domain, data i presentation. Les entitats principals del domini "
    "Flutter són: Excursion, Equipamiento, Solicitud, Reserva i CategoriaActividad."
)

add_image_placeholder(doc, "DIAGRAMA DE CLASSES: entitats del domini Flutter (Excursion, Equipamiento, Solicitud, Reserva, LineaReserva, CategoriaActividad, EstadoExcursion, EstadoEquipamiento...)", height_cm=8)

add_body(doc, "Casos d'ús principals per rol:")
add_bullet(doc, "INVITADO: consultar catàleg d'excursions i material.")
add_bullet(doc, "USUARIO: iniciar sessió, sol·licitar excursions, fer reserves de material, veure el seu historial.")
add_bullet(doc, "EXPERTO: gestionar les sol·licituds assignades, marcar estats.")
add_bullet(doc, "ADMIN: CRUD complet d'excursions, material, usuaris i reservas. Assignar experts.")
add_bullet(doc, "SUPERADMIN: tot el que fa l'ADMIN més la gestió d'admins i la configuració global.")

add_image_placeholder(doc, "DIAGRAMA DE CASOS D'ÚS: actors (Invitat, Usuari, Expert, Admin, Superadmin) amb les seves accions principals", height_cm=8)

# 2.3
add_heading(doc, "2.3 Disseny", level=2, color=COLOR_DARK_GREEN)

add_heading(doc, "2.3.1 Mockups", level=3, color=COLOR_SECONDARY)
add_body(doc,
    "El disseny de l'aplicació s'ha fet pensant en una UI clara i intuïtiva. "
    "Les pantalles principals són:"
)
add_bullet(doc, "Pantalla de login (email + password).")
add_bullet(doc, "Home client: llista d'excursions recomanades i accés ràpid a reserves.")
add_bullet(doc, "Home admin: resum de sol·licituds pendents i alertes d'estoc.")
add_bullet(doc, "Catàleg d'excursions: llista filtrable per categoria i data.")
add_bullet(doc, "Detall d'excursió: informació, participants, material recomanat i botó de sol·licitud.")
add_bullet(doc, "Catàleg de material: llista amb filtre per estat i categoria.")
add_bullet(doc, "Sol·licituds: llista i detall de sol·licituds amb canvi d'estat.")
add_bullet(doc, "Reserves: gestió de reserves de material.")
add_bullet(doc, "Usuaris (admin): llista, creació i edició d'usuaris.")
add_bullet(doc, "Calendari: vista mensual amb totes les excursions.")

add_image_placeholder(doc, "MOCKUPS: pantalles de login, home client, home admin i catàleg d'excursions", height_cm=10)
add_image_placeholder(doc, "MOCKUPS: detall d'excursió, catàleg de material, sol·licituds i reserves", height_cm=10)
add_image_placeholder(doc, "MOCKUPS: formularis de creació (excursió, material, usuari) i pantalla de calendari", height_cm=10)

add_heading(doc, "2.3.2 Paleta de colors", level=3, color=COLOR_SECONDARY)
add_body(doc,
    "La paleta de colors d'Outventura reflecteix la natura i l'aventura. "
    "S'ha definit un tema clar i un tema fosc, tots dos coherents i accessibles."
)

add_image_placeholder(doc, "PALETA DE COLORS: mostrar cada color amb el seu nom, codi hex i exemple d'ús", height_cm=6)

color_headers = ["Nom", "Hex (clar)", "Hex (fosc)", "Ús principal"]
color_rows = [
    ["Primary",          "#588C23", "#7BAF46", "Botons, elements actius, accents"],
    ["PrimaryContainer", "#AFD68D", "#587936", "Fons de cards, xips actius"],
    ["SurfaceContainer", "#3B593F", "#6FAF77", "AppBar, menú lateral, headers"],
    ["Secondary",        "#A6774E", "#D69861", "Accents secundaris, preus, avisos"],
    ["Tertiary",         "#324756", "#4B7491", "Informació addicional, icones info"],
    ["Surface",          "#FFFFFF", "#201814", "Fons general de la pantalla"],
    ["Error",            "#B53A3A", "#EC5E5E", "Errors, alertes, estats crítics"],
]
add_colored_table(doc, color_headers, color_rows)

add_heading(doc, "2.3.3 Usabilitat", level=3, color=COLOR_SECONDARY)
add_body(doc,
    "L'aplicació segueix les guies de Material Design 3 de Google, adaptades "
    "al tema personalitzat d'Outventura. Alguns principis aplicats:"
)
add_bullet(doc, "Navegació consistent: barra de navegació inferior amb les seccions principals.")
add_bullet(doc, "Feedback visual: indicadors de càrrega (loading spinners) i missatges d'error clars.")
add_bullet(doc, "Filtres i cerca: xips de selecció per filtrar per categoria, estat i data sense necessitat de formularis complexos.")
add_bullet(doc, "Formularis validats: tots els camps mostren errors en temps real amb missatges comprensibles.")
add_bullet(doc, "Tema fosc i clar: l'app s'adapta automàticament a la preferència del sistema.")
add_bullet(doc, "Accessibilitat: contrast de colors suficient i etiquetes semàntiques en tots els elements interactius.")

# 2.4
add_heading(doc, "2.4 Tecnologies", level=2, color=COLOR_DARK_GREEN)

add_heading(doc, "2.4.1 Desplegament", level=3, color=COLOR_SECONDARY)
add_body(doc,
    "L'entorn de desenvolupament s'aixeca amb Docker Compose. El fitxer "
    "docker-compose.yml defineix dos serveis:"
)
add_bullet(doc, "postgres: imatge oficial de PostgreSQL 16, amb les credencials de l'entorn de dev.")
add_bullet(doc, "pgadmin: interfície web per a la gestió visual de la base de dades (port 8080).")
add_body(doc,
    "Amb una sola comanda (docker compose up -d) tot l'entorn de base de dades "
    "queda llest. El backend NestJS s'executa a part amb npm run start:dev, "
    "i Prisma gestiona les migracions i el seed inicial."
)
add_body(doc,
    "De cara al desplegament en producció (previst), el backend es containeritzarà "
    "completament i es podria allotjar en un servei de cloud com Railway, Render o "
    "un VPS propi. L'app Flutter es distribuirà via Google Play Store i App Store."
)

add_heading(doc, "2.4.2 Backend", level=3, color=COLOR_SECONDARY)
add_body(doc, "Les tecnologies principals del backend són:")

tech_headers = ["Tecnologia", "Versió", "Funció"]
tech_rows_back = [
    ["NestJS",          "^11.0",  "Framework principal del servidor REST"],
    ["Node.js",         "LTS",    "Runtime d'execució de JavaScript/TypeScript"],
    ["TypeScript",      "5.x",    "Tipatge estàtic i millor mantenibilitat del codi"],
    ["Prisma ORM",      "^7.7",   "Accés a la base de dades i gestió de migracions"],
    ["PostgreSQL",      "16",     "Base de dades relacional principal"],
    ["JWT (passport)",  "^11.0",  "Autenticació basada en tokens"],
    ["bcrypt",          "^6.0",   "Xifrat de contrasenyes"],
    ["Swagger",         "^11.3",  "Documentació automàtica de la API REST"],
    ["class-validator", "^0.15",  "Validació dels DTOs i dades d'entrada"],
    ["Docker Compose",  "v2",     "Orquestració de l'entorn de dev (BD + admin)"],
]
add_colored_table(doc, tech_headers, tech_rows_back)

add_body(doc, "Arquitectura del backend:")
add_bullet(doc, "Estructura modular de NestJS: cada recurs (auth, users, activity, equipment...) té el seu propi mòdul amb controller, service i DTOs.")
add_bullet(doc, "PrismaService: servei compartit injectat a tots els mòduls que necessiten accedir a la base de dades.")
add_bullet(doc, "Guards JWT: protecció dels endpoints que requereixen autenticació.")
add_bullet(doc, "Swagger: documentació interactiva accessible a /api.")

add_heading(doc, "2.4.3 Frontend", level=3, color=COLOR_SECONDARY)
add_body(doc, "Les tecnologies principals del frontend són:")

tech_rows_front = [
    ["Flutter",                "SDK ^3.x",  "Framework multiplataforma (Android, iOS, web)"],
    ["Dart",                   "^3.11",     "Llenguatge de programació"],
    ["Riverpod",               "^3.3",      "Gestió d'estat reactiu"],
    ["Dio",                    "^5.9",      "Client HTTP per a les cridades a la API"],
    ["flutter_secure_storage", "^9.2",      "Emmagatzematge segur del token JWT"],
    ["shared_preferences",     "^2.2",      "Preferències locals (tema, últim login)"],
    ["table_calendar",         "^3.1",      "Vista de calendari mensual"],
    ["calendar_view",          "^2.0",      "Vista setmanal/diària del calendari"],
    ["intl",                   "^0.20",     "Formatació de dates i nombres"],
]
add_colored_table(doc, tech_headers, tech_rows_front)

add_body(doc, "Arquitectura del frontend (Clean Architecture):")
add_bullet(doc, "domain/entities: classes pures (Excursion, Equipamiento, Solicitud, Reserva) sense dependències externes.")
add_bullet(doc, "data/fakes: dades simulades per al desenvolupament offline.")
add_bullet(doc, "features/auth: login, gestió del token i estat d'autenticació.")
add_bullet(doc, "features/outventura/presentation: pàgines, controllers i widgets.")
add_bullet(doc, "core/widgets: components reutilitzables (AppFilterChip, AppChoiceChip, AppChipWrap...).")

add_heading(doc, "2.4.4 Disseny", level=3, color=COLOR_SECONDARY)
add_body(doc,
    "El disseny visual s'ha creat seguint Material Design 3. "
    "Els fitxers clau del tema són:"
)
add_bullet(doc, "app_colors.dart: tota la paleta de colors per al tema clar i fosc.")
add_bullet(doc, "app_text_styles.dart: tipografies i mides de text (titleLarge, headlineSmall, bodyMedium...).")
add_bullet(doc, "app_theme.dart: configuració completa del ThemeData de Flutter (AppBar, ColorScheme, TextTheme).")
add_body(doc,
    "Per al disseny dels mockups s'ha fet servir Figma, seguint un sistema de "
    "components consistent amb el codi Flutter final."
)

add_image_placeholder(doc, "CAPTURES DE PANTALLA: vista de l'app en mode clar i mode fosc mostrant la coherència del tema", height_cm=8)

# 2.5
add_heading(doc, "2.5 Planificació", level=2, color=COLOR_DARK_GREEN)
add_body(doc,
    "El projecte s'ha planificat en fases iteratives, combinant el "
    "desenvolupament del backend i del frontend en paral·lel."
)

plan_headers = ["Fase", "Tasques principals", "Durada estimada"]
plan_rows = [
    ["1. Anàlisi i disseny",       "Definició de requisits, ER, mockups, paleta de colors",               "2 setmanes"],
    ["2. Configuració del projecte","Setup NestJS + Prisma + Docker, setup Flutter + Riverpod",            "1 setmana"],
    ["3. Backend: base",           "Mòduls User, Role, Category, Auth (JWT), Swagger",                    "2 setmanes"],
    ["4. Backend: recursos",       "Mòduls Activity, Equipment, EquipmentStatus, DTOs i validacions",     "2 setmanes"],
    ["5. Frontend: base i auth",   "Tema, components globals, login, gestió de sessions",                 "2 setmanes"],
    ["6. Frontend: funcionalitats","Catàleg, detalls, formularis, sol·licituds, reserves, calendari",     "3 setmanes"],
    ["7. Integració",              "Connexió frontend-backend, gestió d'errors, proves",                  "1 setmana"],
    ["8. Poliment i documentació", "Millores UI, tests, memoria del projecte",                            "1 setmana"],
]
add_colored_table(doc, plan_headers, plan_rows)

add_heading(doc, "2.5.1 Diagrama de Gantt", level=3, color=COLOR_SECONDARY)
add_image_placeholder(doc, "DIAGRAMA DE GANTT: planificació temporal de totes les fases del projecte (setmana a setmana)", height_cm=8)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
#  3. IMPLEMENTACIÓ
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "3. Implementació", level=1, color=COLOR_PRIMARY)

# 3.1
add_heading(doc, "3.1 Desplegament", level=2, color=COLOR_DARK_GREEN)
add_body(doc,
    "L'entorn de desenvolupament s'aixeca amb Docker Compose. "
    "El fitxer docker-compose.yml defineix els serveis necessaris:"
)

add_body(doc, "docker-compose.yml (resum):", bold=True)
p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(1)
run = p.add_run(
    "services:\n"
    "  postgres:\n"
    "    image: postgres:16\n"
    "    environment: { POSTGRES_DB: outventura_db, ... }\n"
    "    ports: [ '5432:5432' ]\n"
    "  pgadmin:\n"
    "    image: dpage/pgadmin4\n"
    "    ports: [ '8080:80' ]"
)
run.font.name = "Courier New"
run.font.size = Pt(9)
run.font.color.rgb = COLOR_TERTIARY

add_body(doc,
    "Per iniciar l'entorn: docker compose up -d. "
    "El backend s'engega amb npm run start:dev (port 3000). "
    "La documentació Swagger és accessible a http://localhost:3000/api."
)
add_body(doc,
    "El seed inicial (prisma/seed.ts) pobla la base de dades amb els rols del "
    "sistema, categories i estats d'equipament. Per executar-lo: npx prisma db seed."
)

add_image_placeholder(doc, "CAPTURA: Docker Desktop amb els contenidors postgres i pgadmin en execució", height_cm=4)
add_image_placeholder(doc, "CAPTURA: interfície Swagger de la API amb tots els endpoints documentats", height_cm=6)

# 3.2
add_heading(doc, "3.2 Backend", level=2, color=COLOR_DARK_GREEN)
add_body(doc,
    "El backend s'ha construït amb NestJS seguint la seva estructura modular. "
    "Cada recurs del domini té el seu propi mòdul independent."
)

add_body(doc, "Mòduls implementats:", bold=True)
modules_headers = ["Mòdul", "Estat", "Endpoints implementats"]
modules_rows = [
    ["Auth",             "✓ Complet",     "POST /auth/login"],
    ["Users",            "En curs",        "GET, POST, PUT, DELETE /users"],
    ["Roles",            "✓ Complet",     "GET /roles (referència)"],
    ["Categories",       "✓ Complet",     "CRUD /category"],
    ["Activities",       "✓ Complet",     "CRUD /activity + assignar categories"],
    ["Equipment",        "✓ Complet",     "CRUD /equipment + assignar categories"],
    ["EquipmentStatus",  "✓ Complet",     "CRUD /equipment-status"],
    ["Solicituds",       "Pendent",        "Previst: CRUD /requests + canvi d'estat"],
    ["Reserves",         "Pendent",        "Previst: CRUD /reservations"],
    ["Notificacions",    "Pendent",        "Previst: push notifications"],
]
add_colored_table(doc, modules_headers, modules_rows)

add_body(doc, "Autenticació i seguretat:", bold=True)
add_bullet(doc, "Les contrasenyes es xifren amb bcrypt (factor de cost per defecte).")
add_bullet(doc, "Els tokens JWT es generen amb @nestjs/jwt i es validen amb passport-jwt.")
add_bullet(doc, "La validació de les dades d'entrada es fa amb class-validator als DTOs.")
add_bullet(doc, "CORS activat per permetre les peticions des del frontend Flutter.")

add_body(doc, "Exemple de flux d'autenticació:", bold=True)
add_body(doc, "1. L'usuari envia email + password a POST /auth/login.")
add_body(doc, "2. El backend valida les credencials i retorna un accessToken JWT i un refreshToken.")
add_body(doc, "3. El frontend desa el token de forma segura amb flutter_secure_storage.")
add_body(doc, "4. Totes les peticions posteriors inclouen la capçalera Authorization: Bearer <token>.")

add_image_placeholder(doc, "CAPTURA: codi del AuthService (login) mostrant la validació i generació del JWT", height_cm=5)
add_image_placeholder(doc, "CAPTURA: exemple de petició a Postman o Swagger mostrant el login i la resposta amb token", height_cm=5)

add_body(doc, "Estructura del projecte backend:", bold=True)
add_body(doc,
    "src/auth/ · src/user/ · src/role/ · src/activity/ · src/equipment/ · "
    "src/equipment-status/ · src/category/ · src/prisma/"
)
add_body(doc,
    "Cada mòdul segueix la mateixa estructura: *.module.ts, *.controller.ts, "
    "*.service.ts i dto/ amb els DTOs de creació i actualització."
)

# 3.3
add_heading(doc, "3.3 Frontend", level=2, color=COLOR_DARK_GREEN)
add_body(doc,
    "El frontend és una aplicació Flutter multiplataforma (Android i iOS principalment). "
    "Utilitza Riverpod per a la gestió d'estat i Dio per a les peticions HTTP a la API."
)

add_body(doc, "Pantalles implementades:", bold=True)
screens_headers = ["Pantalla", "Rol", "Descripció"]
screens_rows = [
    ["Login",                "Tots",          "Formulari d'accés amb email i contrasenya"],
    ["Home Client",          "USUARIO",       "Llista d'excursions i accés ràpid a reserves"],
    ["Home Admin",           "ADMIN/SUPERADMIN", "Resum de sol·licituds pendents i alertes"],
    ["Catàleg excursions",   "Tots",          "Llista filtrable per categoria, data i estat"],
    ["Detall excursió",      "Tots",          "Info completa + material recomanat + botó sol·licitud"],
    ["Formulari excursió",   "ADMIN",         "Creació i edició d'excursions"],
    ["Catàleg material",     "Tots",          "Llista filtrable d'equipament amb estat i preu"],
    ["Formulari material",   "ADMIN",         "Creació i edició de material"],
    ["Sol·licituds",         "Tots",          "Llista de sol·licituds filtrades per rol"],
    ["Detall sol·licitud",   "Tots",          "Vista completa + canvi d'estat + assignació expert"],
    ["Formulari sol·licitud","USUARIO",       "Crear nova sol·licitud d'excursió"],
    ["Reserves",             "ADMIN/USUARIO", "Gestió de reserves de material"],
    ["Detall reserva",       "Tots",          "Línies de reserva i gestió de danys"],
    ["Formulari reserva",    "USUARIO",       "Crear nova reserva de material"],
    ["Usuaris",              "ADMIN",         "Llista, creació i edició d'usuaris"],
    ["Calendari",            "Tots",          "Vista mensual amb totes les excursions"],
    ["Preferències",         "Tots",          "Canvi de tema (clar/fosc), idioma"],
]
add_colored_table(doc, screens_headers, screens_rows)

add_image_placeholder(doc, "CAPTURES: pantalles de home client i home admin en dispositiu Android real o simulador", height_cm=8)
add_image_placeholder(doc, "CAPTURES: catàleg d'excursions amb filtres actius i detall d'una excursió", height_cm=8)
add_image_placeholder(doc, "CAPTURES: catàleg de material i formulari de creació d'equipament", height_cm=8)
add_image_placeholder(doc, "CAPTURES: sol·licituds, detall i formulari de nova sol·licitud", height_cm=8)
add_image_placeholder(doc, "CAPTURES: reserves, detall amb línies i calendari mensual", height_cm=8)

add_body(doc, "Gestió d'estat amb Riverpod:", bold=True)
add_body(doc,
    "S'utilitzen providers i controllers per gestionar l'estat de cada pantalla. "
    "Els controllers encapsulen la lògica de negoci (filtres, canvis d'estat) i "
    "els providers exposen l'estat als widgets de forma reactiva."
)

add_body(doc, "Exemple de selecció de categories amb xips (AppFilterChip):", bold=True)
p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(1)
run = p.add_run(
    "// Selecció múltiple\n"
    "AppChipWrap(\n"
    "  children: CategoriaActividad.values.map((cat) {\n"
    "    final selected = _controller.categorias.contains(cat);\n"
    "    return AppFilterChip(\n"
    "      label: cat.label,\n"
    "      selected: selected,\n"
    "      onSelected: (_) => setState(() =>\n"
    "        _controller.toggleCategoria(cat)),\n"
    "    );\n"
    "  }).toList(),\n"
    ")"
)
run.font.name = "Courier New"
run.font.size = Pt(9)
run.font.color.rgb = COLOR_TERTIARY

add_body(doc,
    "Seguint el mateix patró, AppChoiceChip s'usa per a la selecció individual "
    "(per exemple, l'estat d'un material)."
)

# 3.4
add_heading(doc, "3.4 Disseny", level=2, color=COLOR_DARK_GREEN)
add_body(doc,
    "El disseny visual s'ha implementat amb un sistema de tema personalitzat "
    "que s'aplica globalment a tota l'aplicació. El punt d'entrada del tema és "
    "AppTheme.light i AppTheme.dark, que es passen al MaterialApp."
)
add_body(doc,
    "S'ha posat especial cura en la coherència visual: totes les pantalles fan "
    "servir els mateixos colors, tipografies i components. Els widgets reutilitzables "
    "(AppFilterChip, AppChoiceChip, AppChipWrap) garanteixen que l'estil sigui "
    "consistent a tota l'app."
)

add_image_placeholder(doc, "CAPTURES: comparativa tema clar vs tema fosc en diverses pantalles", height_cm=8)
add_image_placeholder(doc, "CAPTURA: fitxer app_colors.dart o el ColorScheme definit mostrant tots els colors", height_cm=5)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
#  4. MILLORES FUTURES
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "4. Millores futures", level=1, color=COLOR_PRIMARY)
add_body(doc,
    "L'aplicació té un bon punt de partida, però hi ha moltes funcionalitats "
    "que queden pendents per a futures versions:"
)

add_body(doc, "Backend:", bold=True)
add_bullet(doc, "Completar els mòduls de sol·licituds i reserves amb tota la lògica de negoci.")
add_bullet(doc, "Implementar l'endpoint de registro públic (POST /auth/register).")
add_bullet(doc, "Afegir restabliment de contrasenya per email (forgot/reset password).")
add_bullet(doc, "Sistema de notificacions push per avisar de canvis d'estat.")
add_bullet(doc, "Pujada d'imatges per a excursions i material (integració amb S3 o similar).")
add_bullet(doc, "Paginació i filtres avançats a tots els endpoints de llistat.")
add_bullet(doc, "Tests unitaris i d'integració complets.")

add_body(doc, "Frontend:", bold=True)
add_bullet(doc, "Substitució de les dades fake per cridades reals a la API.")
add_bullet(doc, "Pantalla de perfil d'usuari amb edició de dades i canvi de contrasenya.")
add_bullet(doc, "Sistema de notificacions in-app.")
add_bullet(doc, "Mapa integrat per a la ubicació de punts d'inici/fi d'excursions.")
add_bullet(doc, "Valoracions i comentaris d'excursions per part dels usuaris.")
add_bullet(doc, "Exportació de reserves i sol·licituds a PDF.")

add_body(doc, "Negoci:", bold=True)
add_bullet(doc, "Integració amb passarel·la de pagament (Stripe o similar).")
add_bullet(doc, "Estadístiques i panell de control per a administradors.")
add_bullet(doc, "Versió web per a la gestió d'admin des d'escriptori.")
add_bullet(doc, "Internacionalització completa (ca, es, en).")

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
#  5. CONCLUSIONS
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "5. Conclusions", level=1, color=COLOR_PRIMARY)

add_heading(doc, "5.1 Personals", level=2, color=COLOR_DARK_GREEN)
add_body(doc,
    "Aquest projecte ha estat un repte important. Ha estat la primera vegada "
    "que he desenvolupat una aplicació completa des de zero, amb backend i frontend "
    "alhora, i he hagut d'aprendre i aplicar moltes tecnologies noves."
)
add_body(doc,
    "Flutter ha sigut una gran sorpresa: és molt potent i permet fer UIs boniques "
    "amb menys codi del que esperava. Riverpod costa d'entendre al principi, però "
    "un cop l'agafes, la gestió d'estat és molt neta i escalable."
)
add_body(doc,
    "Del backend, el que més m'ha agradat ha sigut NestJS pel seu sistema de mòduls "
    "i la integració amb Swagger, que fa molt fàcil documentar i provar la API. "
    "Prisma també és molt còmode per treballar amb bases de dades sense escriure SQL."
)
add_body(doc,
    "Gestionar el temps ha sigut el punt més difícil. Hi ha moltes funcionalitats "
    "que m'hagués agradat tenir completament acabades, però he après que planificar "
    "bé des del principi és tan important com el codi en si."
)

add_heading(doc, "5.2 Tècniques", level=2, color=COLOR_DARK_GREEN)
add_body(doc,
    "Tècnicament, el projecte ha assolit els objectius principals: "
    "una API REST funcional, una aplicació Flutter amb navegació i UI completes, "
    "un sistema d'autenticació JWT i un entorn de desplegament amb Docker."
)
add_body(doc,
    "He après a estructurar un projecte real amb Clean Architecture al frontend, "
    "cosa que fa el codi molt més mantenible. Al backend, la separació per mòduls "
    "de NestJS ajuda molt quan el projecte creix."
)
add_body(doc,
    "Queda pendent la integració real entre frontend i backend (ara el frontend "
    "usa dades fake), que seria el següent pas lògic. La base tècnica, però, "
    "és sòlida per a fer-ho."
)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
#  6. REFERÈNCIES BIBLIOGRÀFIQUES
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "6. Referències bibliogràfiques", level=1, color=COLOR_PRIMARY)

refs = [
    "Flutter Documentation. (2024). Flutter - Build apps for any screen. https://flutter.dev/docs",
    "Dart Documentation. (2024). Dart programming language. https://dart.dev/docs",
    "Riverpod Documentation. (2024). Riverpod - A reactive state-management library. https://riverpod.dev",
    "NestJS Documentation. (2024). NestJS - A progressive Node.js framework. https://docs.nestjs.com",
    "Prisma Documentation. (2024). Prisma - Next-generation ORM. https://www.prisma.io/docs",
    "PostgreSQL Documentation. (2024). PostgreSQL: The world's most advanced open source database. https://www.postgresql.org/docs/",
    "Material Design 3. (2024). Material Design guidelines. https://m3.material.io",
    "Docker Documentation. (2024). Docker - Containerization platform. https://docs.docker.com",
    "JWT.io. (2024). JSON Web Tokens. https://jwt.io/introduction",
    "Swagger / OpenAPI. (2024). OpenAPI Specification. https://swagger.io/specification/",
]
for i, ref in enumerate(refs, 1):
    p = doc.add_paragraph(f"[{i}] {ref}")
    p.paragraph_format.space_after = Pt(4)
    for run in p.runs:
        run.font.size = Pt(10)

# ── Save ───────────────────────────────────────────────────────────────────────
output_path = "memoria_outventura.docx"
doc.save(output_path)
print(f"Document generated: {output_path}")
