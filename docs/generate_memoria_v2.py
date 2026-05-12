"""
Outventura — Memoria del Projecte (v2)
Generates memoria_outventura_v2.docx
Typography: Montserrat (headings) + Open Sans (body)
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Colors ─────────────────────────────────────────────────────────────────────
PRIMARY         = RGBColor(0x58, 0x8C, 0x23)
DARK_GREEN      = RGBColor(0x3B, 0x59, 0x3F)
LIGHT_GREEN     = RGBColor(0xAF, 0xD6, 0x8D)
SECONDARY       = RGBColor(0xA6, 0x77, 0x4E)
TERTIARY        = RGBColor(0x32, 0x47, 0x56)
WHITE           = RGBColor(0xFF, 0xFF, 0xFF)
DARK            = RGBColor(0x20, 0x18, 0x14)
GRAY            = RGBColor(0x55, 0x55, 0x55)
LIGHT_GRAY      = RGBColor(0x99, 0x99, 0x99)

# Fonts
FONT_HEADING = "Montserrat"
FONT_BODY    = "Open Sans"

doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(2.5)


# ── Helpers ────────────────────────────────────────────────────────────────────

def _hex(color):
    return str(color).upper()

def _cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def heading1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(text)
    run.font.name = FONT_HEADING
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = PRIMARY
    return p

def heading2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.name = FONT_HEADING
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = DARK_GREEN
    return p

def heading3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = FONT_HEADING
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = SECONDARY
    return p

def body(text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = FONT_BODY
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    if level > 0:
        p.paragraph_format.left_indent = Cm(1.5 * level)
    for run in p.runs:
        run.font.name = FONT_BODY
        run.font.size = Pt(11)
    # Clear default and add our own
    p.clear()
    run = p.add_run(text)
    run.font.name = FONT_BODY
    run.font.size = Pt(11)
    return p

def code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = TERTIARY
    return p

def img_placeholder(desc, h=7):
    """Gray box indicating where an image should go."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    cell.width = Cm(14)
    _cell_bg(cell, 'F0F0F0')
    trPr = tbl.rows[0]._tr.get_or_add_trPr()
    trH = OxmlElement('w:trHeight')
    trH.set(qn('w:val'), str(int(h * 567)))
    trH.set(qn('w:hRule'), 'atLeast')
    trPr.append(trH)
    cp = cell.paragraphs[0]
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cp.add_run(f"\n\n[ IMAGEN: {desc} ]\n\n")
    run.font.name = FONT_BODY
    run.font.size = Pt(10)
    run.font.color.rgb = LIGHT_GRAY
    run.italic = True
    doc.add_paragraph()

def table(headers, rows, hdr_color=None):
    hc = hdr_color or DARK_GREEN
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ''
        _cell_bg(c, _hex(hc))
        r = c.paragraphs[0].add_run(h)
        r.font.name = FONT_BODY
        r.font.size = Pt(9)
        r.font.bold = True
        r.font.color.rgb = WHITE
    for ri, rd in enumerate(rows):
        for ci, val in enumerate(rd):
            c = t.rows[ri+1].cells[ci]
            c.text = ''
            r = c.paragraphs[0].add_run(str(val))
            r.font.name = FONT_BODY
            r.font.size = Pt(9)
    doc.add_paragraph()
    return t

def page_break():
    doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
#  PORTADA
# ══════════════════════════════════════════════════════════════════════════════

for _ in range(2):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Manuel Navalon Fornes")
run.font.name = FONT_HEADING
run.font.size = Pt(14)
run.font.color.rgb = DARK_GREEN

doc.add_paragraph()
img_placeholder("LOGO OUTVENTURA (icono de muntanya amb text)", h=5)
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("OUTVENTURA")
run.font.name = FONT_HEADING
run.font.size = Pt(36)
run.font.bold = True
run.font.color.rgb = PRIMARY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("L'aventura al teu abast")
run.font.name = FONT_BODY
run.font.size = Pt(14)
run.font.italic = True
run.font.color.rgb = SECONDARY

for _ in range(3):
    doc.add_paragraph()

lines = [
    ("CFGS DAW", True, 13),
    ("Mòdul Professional Projecte Intermodular", False, 12),
    ("Desenvolupament d'Aplicacions Multiplataforma", False, 12),
    ("IES L'ESTACIÓ", True, 13),
    ("Curs 2025-26", False, 11),
]
for txt, bold, sz in lines:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(txt)
    run.font.name = FONT_HEADING
    run.font.size = Pt(sz)
    run.font.bold = bold
    run.font.color.rgb = DARK_GREEN

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Alumne: Manuel Navalon Fornes")
run.font.name = FONT_BODY
run.font.size = Pt(12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Tutor: __________________________")
run.font.name = FONT_BODY
run.font.size = Pt(12)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  ÍNDEX
# ══════════════════════════════════════════════════════════════════════════════
heading1("Índex")
idx = [
    "1.  Introducció",
    "    1.1  Idea de negoci",
    "    1.2  Objectius",
    "    1.3  Anàlisi del mercat i segmentació. Competència. DAFO",
    "    1.4  Alineació amb ODS",
    "2.  Anàlisi dels requeriments",
    "    2.1  Entitat Relació",
    "    2.2  Diagrama de classes i casos d'ús",
    "    2.3  Disseny",
    "        2.3.1  Mockups",
    "        2.3.2  Paleta de colors",
    "        2.3.3  Usabilitat",
    "    2.4  Tecnologies",
    "        2.4.1  Desplegament",
    "        2.4.2  Backend",
    "        2.4.3  Frontend",
    "        2.4.4  Disseny",
    "    2.5  Planificació",
    "        2.5.1  Diagrama de Gantt",
    "3.  Implementació",
    "    3.1  Desplegament",
    "    3.2  Backend",
    "    3.3  Frontend",
    "    3.4  Disseny",
    "4.  Millores futures",
    "5.  Conclusions",
    "    5.1  Personals",
    "    5.2  Tècniques",
    "6.  Referències bibliogràfiques",
]
for e in idx:
    p = doc.add_paragraph(e)
    p.paragraph_format.space_after = Pt(2)
    for r in p.runs:
        r.font.name = FONT_BODY
        r.font.size = Pt(11)

page_break()


# ══════════════════════════════════════════════════════════════════════════════
#  1. INTRODUCCIÓ
# ══════════════════════════════════════════════════════════════════════════════
heading1("1. Introducció")

heading2("1.1 Idea de negoci")
body(
    "Outventura és una aplicació mòbil destinada a empreses i associacions que "
    "organitzen activitats a l'aire lliure: excursions de muntanya, rutes aquàtiques, "
    "sortides de neu i acampades. Neix per donar solució a la gestió manual "
    "(excels, WhatsApp, paper) que encara fan servir moltes d'aquestes entitats."
)
body(
    "L'app ofereix dos grans blocs funcionals: d'una banda, la gestió del catàleg "
    "d'excursions i del material d'aventura per part dels administradors; de l'altra, "
    "la consulta, sol·licitud i reserva per part dels clients. A més, incorpora "
    "la figura de l'expert (guia assignat) que pot gestionar les sol·licituds "
    "que li arriben."
)
body(
    "El model de negoci es basa en: (1) preu per participant en cada excursió, "
    "(2) lloguer diari de material i (3) càrrec per danys en la devolució. Tot "
    "automatitzat i visible dins l'app, tant per a l'admin com per al client."
)

heading2("1.2 Objectius")
body("Objectius funcionals:")
bullet("Gestió completa d'excursions: crear, editar, eliminar, filtrar per categoria/data/estat.")
bullet("Gestió d'equipament: inventari, preus de lloguer, estats (disponible, agotado, mantenimiento, fora de servei).")
bullet("Sistema de sol·licituds: el client sol·licita participar en una excursió, l'admin o expert l'accepta o rebutja.")
bullet("Sistema de reserves de material: lligat a sol·licituds o independent, amb control de danys en la devolució.")
bullet("Calendari: vista mensual amb excursions, reserves i sol·licituds.")
bullet("Autenticació JWT amb rols (SUPERADMIN, ADMIN, EXPERTO, USUARIO, INVITADO).")
body("Objectius tècnics:")
bullet("App multiplataforma (Android + iOS) amb Flutter i Dart.")
bullet("Backend REST amb NestJS, TypeScript, Prisma ORM i PostgreSQL.")
bullet("Documentació de la API amb Swagger (OpenAPI).")
bullet("Entorn containeritzat amb Docker Compose (BD + pgAdmin).")
bullet("Arquitectura neta al frontend (Clean Architecture) i modular al backend.")
bullet("UI coherent amb sistema de tema (clar/fosc) i design system intern.")

heading2("1.3 Anàlisi del mercat i segmentació. Competència. DAFO")
body(
    "El sector del turisme actiu a Espanya supera els 4.500 milions d'euros "
    "anuals i creix un 8% interanual. Tanmateix, la majoria d'empreses petites "
    "encara gestionen les reserves per telèfon o formularis web genèrics."
)
body("Competència directa i indirecta:", bold=True)
table(
    ["Competidor", "Què fa", "Limitació respecte Outventura"],
    [
        ["Civitatis", "Marketplace d'activitats turístiques", "No gestiona material ni equips interns"],
        ["Wikiloc", "Rutes GPS compartides", "No té reserves ni gestió empresarial"],
        ["Fareharbor", "Reserves online per operadors", "Preu alt, no open source, sense inventari"],
        ["Google Forms", "Formularis genèrics", "Zero automatització, sense dashboard"],
    ]
)
body("Segmentació del mercat:", bold=True)
bullet("Empreses d'activitats (escalada, kayak, senderisme).")
bullet("Associacions esportives amb magatzem de material.")
bullet("Centres d'esplai, colònies i grups escolars.")
bullet("Clubs de muntanya i guies independents.")

body("Anàlisi DAFO:", bold=True)
table(
    ["", "Positiu", "Negatiu"],
    [
        ["Intern",
         "Fortaleses:\n• Solució integral (gestió + client)\n• UI moderna Material 3\n• Arquitectura escalable i modular\n• Multiplataforma amb un sol codebase",
         "Debilitats:\n• Desenvolupador únic (temps limitat)\n• Backend en construcció parcial\n• Sense integració de pagament real\n• Tests limitats"],
        ["Extern",
         "Oportunitats:\n• Mercat en creixement post-COVID\n• Poca competència directa low-cost\n• Digitalització obligada del sector\n• Possible SaaS multi-tenant",
         "Amenaces:\n• Apps genèriques de gestió\n• Cost d'adquisició de clients B2B\n• Dependència de stores (Apple/Google)\n• Canvis en APIs de tercers"],
    ]
)

heading2("1.4 Alineació amb objectius, valors, sostenibilitat i futur (ODS)")
body("El projecte s'alinea amb els següents Objectius de Desenvolupament Sostenible:")
bullet("ODS 3 — Salut i benestar: promou l'activitat física i la connexió amb la natura.")
bullet("ODS 8 — Treball decent: digitalitza i professionalitza el sector, creant valor per a guies i petites empreses.")
bullet("ODS 11 — Comunitats sostenibles: fomenta el turisme local i de proximitat.")
bullet("ODS 12 — Consum responsable: el control d'inventari evita compres innecessàries i allarga la vida del material.")
bullet("ODS 13 — Acció climàtica: redueix el paper i promou activitats de baix impacte ambiental.")
body(
    "Valors de marca: accessibilitat (activitats per a tots els nivells), "
    "seguretat (control de material i guies qualificats), respecte ambiental "
    "i transparència (preus clars, sense sorpreses)."
)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  2. ANÀLISI DELS REQUERIMENTS
# ══════════════════════════════════════════════════════════════════════════════
heading1("2. Anàlisi dels requeriments")

heading2("2.1 Entitat Relació")
body(
    "La base de dades relacional (PostgreSQL 16) es gestiona amb Prisma ORM. "
    "S'ha dissenyat amb 6 entitats principals:"
)
table(
    ["Entitat", "PK", "Camps principals", "Relacions"],
    [
        ["Role", "id_role", "code (UNIQUE), description", "1:N → User"],
        ["User", "id_user", "name, surname, email (UNIQUE), phone, photo, password, status, timestamps", "N:1 → Role"],
        ["Category", "id_category", "code (UNIQUE), description", "M:N → Equipment, M:N → Activity"],
        ["EquipmentStatus", "id_status", "code (UNIQUE), description", "1:N → Equipment"],
        ["Equipment", "id_equipment", "title, description, price_per_day, units", "N:1 → EquipmentStatus, M:N → Category"],
        ["Activity", "id_activity", "title, description, init_date, end_date, difficulty, max_participants, start_end_point", "M:N → Category"],
    ]
)
body(
    "Les relacions molts-a-molts (Category ↔ Equipment, Category ↔ Activity) es "
    "gestionen amb taules intermèdies implícites de Prisma. El camp difficulty "
    "d'Activity és un enter de 0 a 100 que representa el percentatge de dificultat."
)

img_placeholder("DIAGRAMA E/R complet: Role, User, Category, Activity, Equipment, EquipmentStatus amb cardinalitats, PKs i FKs", h=10)

heading2("2.2 Diagrama de classes i casos d'ús")
body("Entitats del domini Flutter (capa domain/entities):", bold=True)
table(
    ["Classe", "Atributs principals", "Enums associats"],
    [
        ["Excursion", "id, puntoInicio, puntoFin, imagenAsset, fechaInicio/Fin, categorias, numeroParticipantes, descripcion, precio, materialesPorParticipante", "EstadoExcursion (disponible, pendiente, confirmada, enCurso, finalizada, cancelada)"],
        ["Equipamiento", "id, nombre, descripcion, categorias, stock, stockTotal, estado, precioAlquilerDiario, cargoPorDanio, imagenAsset", "EstadoEquipamiento (disponible, agotado, mantenimiento, fueraDeServicio)"],
        ["Solicitud", "id, idExcursion, numeroParticipantes, estado, idExperto, idUsuario, idReserva, materialesSolicitados, precioTotal", "EstadoSolicitud (pendiente, confirmada, enCurso, finalizada, cancelada)"],
        ["Reserva", "id, idUsuario, lineas (List<LineaReserva>), idExcursion, fechaInicio/Fin, estado, cargoDanios, itemsDaniados", "EstadoReserva (pendiente, confirmada, enCurso, finalizada, cancelada)"],
        ["Usuario", "id, nombre, apellidos, email, telefono, rol, foto, activo", "TipoRol (superadmin, admin, experto, usuario, invitado)"],
        ["CategoriaActividad", "—", "acuatico, nieve, montana, camping"],
    ]
)

img_placeholder("DIAGRAMA DE CLASSES UML: totes les entitats del domini Flutter amb atributs, mètodes factory i relacions", h=9)

body("Casos d'ús per actor:", bold=True)
table(
    ["Actor", "Casos d'ús principals"],
    [
        ["Invitat", "Consultar catàleg d'excursions · Consultar catàleg de material"],
        ["Usuari", "Login · Sol·licitar excursió · Reservar material · Veure historial · Editar perfil · Canviar contrasenya"],
        ["Expert", "Gestionar sol·licituds assignades · Canviar estat de sol·licitud · Veure calendari"],
        ["Admin", "CRUD excursions · CRUD material · CRUD usuaris · Gestionar reserves · Gestionar sol·licituds · Assignar experts · Aprovar/rebutjar reserves · Registrar devolucions"],
        ["Superadmin", "Tot l'anterior + Gestionar admins + Configuració global del sistema"],
    ]
)

img_placeholder("DIAGRAMA DE CASOS D'ÚS UML: 5 actors amb les seves accions dins del sistema", h=9)

heading2("2.3 Disseny")
heading3("2.3.1 Mockups")
body(
    "L'app compta amb 17 pantalles principals, dissenyades amb Material Design 3 "
    "i adaptades a la identitat visual d'Outventura. Les pantalles es divideixen "
    "per flux d'usuari:"
)
body("Flux d'autenticació:", bold=True)
bullet("Login: fons amb imatge de paisatge (Camino.jpg), card semitransparent amb email + password, toggle de visibilitat, enllaços a registre i forgot password.")
bullet("Perfil: edició de dades personals + canvi de contrasenya amb validació en temps real.")

body("Flux de client:", bold=True)
bullet("Home Client: salutació personalitzada, estadístiques (reserves, sol·licituds, pendents), sol·licituds recents i noves excursions.")
bullet("Sol·licitar excursió: selecció d'excursió, nombre de participants, recalcul automàtic de materials recomanats, càlcul de preu total.")
bullet("Reservar material: selecció de dates, línies de material (equipament + quantitat), preu total.")

body("Flux d'administració:", bold=True)
bullet("Home Admin: estadístiques globals, botons de gestió ràpida, sol·licituds pendents.")
bullet("CRUD Excursions: formulari amb punt inici/fi, dates, hores, dificultad, participants, preu, categories (múltiples), imatge.")
bullet("CRUD Material: formulari amb nom, descripció, stock, preu lloguer, càrrec per dany, categories, estat, imatge.")
bullet("CRUD Usuaris: formulari amb dades personals, rol (dropdown), estat actiu/inactiu.")
bullet("Gestió Reserves: aprovar, rebutjar, cancelar, registrar devolució (amb selecció d'ítems danyats).")
bullet("Gestió Sol·licituds: acceptar, rebutjar, assignar expert.")

body("Vista compartida:", bold=True)
bullet("Calendari: vista mensual amb badges 'R' (reserves) i 'S' (sol·licituds), llistat d'events del dia seleccionat.")

img_placeholder("MOCKUPS: pantalla de Login (amb fons Camino.jpg) + Home Client + Home Admin", h=10)
img_placeholder("MOCKUPS: Catàleg d'excursions amb filtres + Detall excursió + Formulari excursió", h=10)
img_placeholder("MOCKUPS: Catàleg de material + Formulari material + Vista calendari", h=10)
img_placeholder("MOCKUPS: Sol·licituds + Detall sol·licitud + Formulari sol·licitud", h=10)
img_placeholder("MOCKUPS: Reserves + Detall reserva (amb línies) + Formulari reserva", h=10)
img_placeholder("MOCKUPS: Gestió d'usuaris + Formulari usuari + Perfil personal", h=10)

heading3("2.3.2 Paleta de colors")
body(
    "La paleta està inspirada en la natura i l'aventura. S'han definit dos temes "
    "complets (clar i fosc) amb suficient contrast per complir WCAG AA."
)
table(
    ["Nom", "Clar (hex)", "Fosc (hex)", "Ús"],
    [
        ["Primary",             "#588C23", "#7BAF46", "Botons principals, accents, FABs"],
        ["Primary Container",   "#AFD68D", "#587936", "Fons de cards actives, xips seleccionats"],
        ["Surface Container",   "#3B593F", "#6FAF77", "AppBar, headers, menú lateral"],
        ["Secondary",           "#A6774E", "#D69861", "Preus, avisos, accents secundaris"],
        ["Tertiary",            "#324756", "#4B7491", "Informació contextual, codi, icones info"],
        ["On Tertiary",         "#88C1E9", "#1A2A3D", "Text sobre fons teriari"],
        ["Surface",             "#FFFFFF", "#201814", "Fons general de l'app"],
        ["On Surface",          "#201814", "#FFFFFF", "Text principal"],
        ["On Primary",          "#F2F0F2", "#171512", "Text sobre botons primaris"],
        ["On Surface Variant",  "#9A979A", "#AFA9AF", "Text secundari, hints"],
        ["Error",               "#B53A3A", "#EC5E5E", "Errors, eliminar, cancel·lar"],
    ]
)

img_placeholder("PALETA VISUAL: quadrats de cada color amb el seu hex, tant en tema clar com fosc, en format de targeta", h=6)

heading3("2.3.3 Usabilitat")
body("Principis d'usabilitat aplicats:")
bullet("Navegació inferior persistent (4 pestañes: Inici, Excursions, Equipament, Calendari) que dóna accés ràpid a les seccions principals sense scroll.")
bullet("Filtres amb xips (FilterChip / ChoiceChip): els filtres es mostren en un bottom sheet amb drag handle, permetent filtrar per categoria, estat i rang de dates sense pantalla nova.")
bullet("Feedback immediat: spinners durant les càrregues (simulades amb 800ms de delay), missatges d'error en camp, diàlegs de confirmació abans d'accions destructives.")
bullet("Formularis validats en temps real: validators centralitzats (campoObligatorio, email regex, enteroPositivo, decimalPositivo, longitudMinima).")
bullet("Adaptat al teclat: SingleChildScrollView en tots els formularis per evitar que el teclat tapi camps.")
bullet("Tema automàtic: l'app detecta la preferència del sistema (clar/fosc) i permet canvi manual des de preferències.")
bullet("Design system intern: catàleg de components (catalog/) que documenta tots els widgets reutilitzables com un storybook.")
bullet("Accessibilitat: contrast mínim 4.5:1 en text, etiquetes semàntiques, icones amb tooltip.")

heading2("2.4 Tecnologies")
heading3("2.4.1 Desplegament")
body("L'entorn de desenvolupament:")
table(
    ["Component", "Tecnologia", "Configuració"],
    [
        ["Base de dades", "PostgreSQL 16 (Docker)", "Container outventura_postgres, port 5432, vol. persistent"],
        ["Admin BD", "pgAdmin 4 (Docker)", "Container outventura_pgadmin, port 8080"],
        ["Orquestració", "Docker Compose v2", "docker-compose.yml amb 2 serveis + volume"],
        ["Backend runtime", "Node.js LTS + NestJS", "npm run start:dev (watch mode, port 3000)"],
        ["Migracions", "Prisma Migrate + Seed", "npx prisma migrate dev / npx prisma db seed"],
        ["Frontend runtime", "Flutter SDK ^3.x", "flutter run (Android/iOS/web)"],
    ]
)
body(
    "De cara a producció (previst): backend en Docker complet amb Dockerfile + "
    "orquestració multi-container. Possibles plataformes: Railway, Render, VPS. "
    "L'app Flutter es publicarà a Google Play Store i Apple App Store."
)

heading3("2.4.2 Backend")
table(
    ["Tecnologia", "Versió", "Funció"],
    [
        ["NestJS",              "^11.0",    "Framework REST modular amb DI"],
        ["TypeScript",          "5.x",      "Tipatge estàtic, interfaces, decoradors"],
        ["Prisma ORM",          "^7.7",     "Mapping objecte-relacional, migracions, seed"],
        ["PostgreSQL",          "16",       "SGBD relacional principal"],
        ["@nestjs/jwt",         "^11.0",    "Generació i validació de tokens JWT"],
        ["passport-jwt",        "^4.0",     "Estratègia d'autenticació"],
        ["bcrypt",              "^6.0",     "Hash de contrasenyes (factor 10)"],
        ["class-validator",     "^0.15",    "Validació declarativa de DTOs"],
        ["class-transformer",   "^0.5",     "Transformació de DTOs"],
        ["@nestjs/swagger",     "^11.3",    "Documentació OpenAPI automàtica"],
        ["@nestjs/config",      "^4.0",     "Variables d'entorn (.env)"],
        ["Docker Compose",      "v2",       "Orquestració PostgreSQL + pgAdmin"],
    ]
)

heading3("2.4.3 Frontend")
table(
    ["Tecnologia", "Versió", "Funció"],
    [
        ["Flutter",                 "SDK ^3.x",  "Framework UI multiplataforma"],
        ["Dart",                    "^3.11",     "Llenguatge principal"],
        ["flutter_riverpod",        "^3.3",      "Gestió d'estat reactiu (providers + notifiers)"],
        ["Dio",                     "^5.9",      "Client HTTP (preparat, no connectat)"],
        ["flutter_secure_storage",  "^9.2",      "Emmagatzematge segur de tokens JWT"],
        ["shared_preferences",      "^2.2",      "Preferències locals (tema, idioma)"],
        ["table_calendar",          "^3.1",      "Calendari mensual interactiu"],
        ["calendar_view",           "^2.0",      "Vista setmanal/diària complementària"],
        ["intl",                    "^0.20",     "Formatació de dates en espanyol"],
        ["uuid",                    "^4.5",      "Generació d'IDs temporals (mode fake)"],
    ]
)

heading3("2.4.4 Disseny")
body("Eines i recursos de disseny:")
bullet("Material Design 3 (M3) com a guia base de components i patrons.")
bullet("Figma per als mockups inicials.")
bullet("Sistema de tema personalitzat a Flutter: app_colors.dart, app_text_styles.dart, app_theme.dart.")
bullet("Tipografia de l'app: pes de 400 a 800, mides de 9pt a 26pt, sense font externa (system font).")
bullet("Catàleg de components (catalog/) com a storybook intern per validar l'estil de cada widget.")

img_placeholder("CAPTURA: catàleg de components (catalog_page.dart) mostrant la demostració de botons, inputs i cards", h=7)

heading2("2.5 Planificació")
table(
    ["Fase", "Durada", "Tasques"],
    [
        ["1. Anàlisi",          "2 setm.", "Requisits, ER, mockups, paleta, plantejament de l'arquitectura"],
        ["2. Setup",            "1 setm.", "NestJS + Prisma + Docker · Flutter + Riverpod · CI bàsic"],
        ["3. Backend base",     "2 setm.", "Mòduls Role, User, Auth (JWT + bcrypt), Category"],
        ["4. Backend recursos", "2 setm.", "Mòduls Activity, Equipment, EquipmentStatus, Swagger complet"],
        ["5. Frontend base",    "2 setm.", "Tema, widgets globals, login, gestió de sessió, navegació"],
        ["6. Frontend features","3 setm.", "Catàleg, formularis, sol·licituds, reserves, calendari, filtres"],
        ["7. Integració",       "1 setm.", "Connexió Dio ↔ API, gestió d'errors, interceptors JWT"],
        ["8. Poliment",         "1 setm.", "Millores UI, tests, memòria, preparació de presentació"],
    ]
)
heading3("2.5.1 Diagrama de Gantt")
img_placeholder("DIAGRAMA DE GANTT: 8 fases distribuïdes en 14 setmanes (abril - juliol 2026)", h=7)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  3. IMPLEMENTACIÓ
# ══════════════════════════════════════════════════════════════════════════════
heading1("3. Implementació")

heading2("3.1 Desplegament")
body("L'entorn de desenvolupament s'organitza així:", bold=True)
code_block(
    "docker-compose.yml\n"
    "├── postgres (PostgreSQL 16)\n"
    "│   ├── POSTGRES_USER: outventura\n"
    "│   ├── POSTGRES_PASSWORD: outventura123\n"
    "│   ├── POSTGRES_DB: outventura_db\n"
    "│   └── port: 5432\n"
    "├── pgadmin (dpage/pgadmin4)\n"
    "│   ├── admin@outventura.com / admin123\n"
    "│   └── port: 8080\n"
    "└── volume: postgres_data (persistent)"
)
body("Comandes principals de l'entorn:")
bullet("docker compose up -d → aixeca BD i pgAdmin.")
bullet("npx prisma migrate dev → aplica migracions pendents.")
bullet("npx prisma db seed → pobla roles (SUPER, ADMIN, USER, GUEST) i 4 usuaris inicials.")
bullet("npm run start:dev → engega el backend en mode watch (port 3000).")
bullet("http://localhost:3000/api → interfície Swagger amb tots els endpoints.")

img_placeholder("CAPTURA: Docker Desktop amb contenidors postgres i pgadmin en marxa", h=4)
img_placeholder("CAPTURA: pgAdmin mostrant les taules del schema Outventura", h=5)
img_placeholder("CAPTURA: Swagger UI amb tots els tags (Auth, Users, Roles, Categories, Activities, Equipment, Equipment Status)", h=6)

heading2("3.2 Backend")
body(
    "El backend segueix l'estructura modular de NestJS. Cada recurs del domini "
    "té el seu propi mòdul amb controller, service i DTOs."
)

body("Estructura de carpetes:", bold=True)
code_block(
    "src/\n"
    "├── main.ts           ← Bootstrap, CORS, ValidationPipe, Swagger\n"
    "├── app.module.ts     ← Importa tots els mòduls\n"
    "├── prisma/           ← PrismaService (global) + PrismaModule\n"
    "├── auth/             ← Login amb JWT + bcrypt\n"
    "├── role/             ← CRUD de rols\n"
    "├── user/             ← CRUD d'usuaris\n"
    "├── category/         ← CRUD de categories\n"
    "├── activity/         ← CRUD d'activitats + assignar categories\n"
    "├── equipment/        ← CRUD d'equipament + assignar categories\n"
    "└── equipment-status/ ← CRUD d'estats de material"
)

body("Mòduls implementats:", bold=True)
table(
    ["Mòdul", "Endpoints", "Operacions DB clau", "Validacions DTO"],
    [
        ["Auth", "POST /auth/login", "findUnique(email) + bcrypt.compare + JWT sign", "email: @IsEmail, password: @MinLength(8)"],
        ["Role", "CRUD /role", "Duplicate code check, count users before delete", "code: @IsNotEmpty, description: optional"],
        ["User", "CRUD /user", "Duplicate email check, role exists check, strips password from response", "name/surname: @IsNotEmpty, email: @IsEmail, password: @MinLength(8), roleId: @IsInt"],
        ["Category", "CRUD /category", "Duplicate code check, includes _count equipments/activities", "code: @MaxLength(50), description: optional"],
        ["Activity", "CRUD /activity + POST :id/category/:catId", "Many-to-many connect categories, includes categories", "title: @IsNotEmpty, dates: @IsDateString, difficulty: @Min(0)@Max(100), max_participants: @Min(1)"],
        ["Equipment", "CRUD /equipment + POST :id/category/:catId", "Many-to-many connect categories, includes status+categories", "title: @IsNotEmpty, price_per_day: @Min(0), units: @Min(1), statusId: @IsInt"],
        ["EquipmentStatus", "CRUD /equipment-status", "Duplicate code check", "code: @MaxLength(20), description: @MaxLength(255)"],
    ]
)

body("Autenticació (flux detallat):", bold=True)
bullet("1. Client envia POST /auth/login amb { email, password }.")
bullet("2. AuthService busca l'usuari per email (inclou role).")
bullet("3. Compara la password amb bcrypt.compare().")
bullet("4. Si és vàlid, genera JWT amb payload { sub: id, email, role: code }. Expiració: 1 dia.")
bullet("5. Retorna { user: { id, name, email, role }, access_token }.")
bullet("6. El frontend desa el token a flutter_secure_storage i l'envia en capçalera Authorization: Bearer <token>.")

body("Seguretat implementada:", bold=True)
bullet("Contrasenyes hashejades amb bcrypt (factor 10) al seed.")
bullet("Validació global amb ValidationPipe (whitelist + forbidNonWhitelisted + transform).")
bullet("CORS habilitat per a desenvolupament.")
bullet("Swagger amb documentació completa (tags, operacions, respostes).")

body("Pendent de completar al backend:", bold=True)
bullet("Guards JWT als endpoints protegits (ara estan oberts).")
bullet("Hash de password al UserService.create() (ara es desa tal qual).")
bullet("Secret JWT des de variable d'entorn (ara hardcoded).")
bullet("Mòduls de Solicitud i Reserva (definits al frontend però no al backend).")
bullet("Refresh token, registro i forgot/reset password.")

img_placeholder("CAPTURA: codi de auth.service.ts mostrant el flux de login i generació JWT", h=5)
img_placeholder("CAPTURA: exemple de petició i resposta al endpoint POST /auth/login (Swagger o Postman)", h=5)
img_placeholder("CAPTURA: user.service.ts mostrant la verificació de duplicats i eliminació del password de la resposta", h=5)

heading2("3.3 Frontend")
body(
    "El frontend és una app Flutter multiplataforma organitzada en features "
    "seguint Clean Architecture. Utilitza Riverpod per a la gestió d'estat reactiu."
)

body("Estructura de carpetes:", bold=True)
code_block(
    "lib/\n"
    "├── main.dart              ← Entry point, ProviderScope, MaterialApp\n"
    "├── app/theme/             ← AppColors, AppTextStyles, AppTheme (light/dark)\n"
    "├── catalog/               ← Design system intern (demos de components)\n"
    "├── core/\n"
    "│   ├── network/           ← ApiDelay, AuthStorage (comentat), DioClient (comentat)\n"
    "│   ├── utils/             ← DateFormatter, FormValidators, IdGenerator\n"
    "│   └── widgets/           ← AddFab, AppButtons, AppChip, AppDateSelector,\n"
    "│                             AppDropdown, AppImagePicker, AppInput, AppTag,\n"
    "│                             AppTimeSelector, ConfirmDialog, DetailSection,\n"
    "│                             FilterBottomSheet, FilterDateRangeRow\n"
    "├── features/\n"
    "│   ├── auth/\n"
    "│   │   ├── domain/entities/   ← Usuario, TipoRol\n"
    "│   │   ├── data/fakes/        ← usuariosFake (4 users)\n"
    "│   │   └── presentation/      ← LoginPage (4 versions), ProfileFormPage,\n"
    "│   │                             LoginController, UserFormController,\n"
    "│   │                             currentUserProvider, usersProvider\n"
    "│   ├── outventura/\n"
    "│   │   ├── domain/entities/   ← Excursion, Equipamiento, Solicitud, Reserva, CategoriaActividad\n"
    "│   │   ├── data/fakes/        ← Dades simulades per a dev offline\n"
    "│   │   ├── services/          ← PricingService, Resolvers\n"
    "│   │   └── presentation/\n"
    "│   │       ├── pages/         ← 12 pantalles + 5 formularis\n"
    "│   │       ├── controllers/   ← 8 controladors de pàgina/formulari\n"
    "│   │       ├── providers/     ← 5 providers (CRUD + filtrats)\n"
    "│   │       └── widgets/       ← Cards, StatCard, etc.\n"
    "│   └── preferences/           ← Gestió de tema clar/fosc\n"
)

body("Pantalles implementades (17 total):", bold=True)
table(
    ["Pantalla", "Fitxer", "Funcionalitat clau"],
    [
        ["Login", "login_page.dart", "Autenticació amb validació, fons amb imatge, navigate a MainScaffold"],
        ["Main Scaffold", "main_scaffold.dart", "BottomNav 4 tabs, detecció de rol per mostrar Home diferent"],
        ["Home Client", "home_client_page.dart", "Stats, salutació, sol·licituds recents, noves excursions"],
        ["Home Admin", "home_admin_page.dart", "Stats globals, botons gestió, sol·licituds pendents"],
        ["Excursions", "excursions_page.dart", "Llistat filtrable, CRUD amb FAB, filtres per estat/categoria/data"],
        ["Equipment", "equipment_page.dart", "Llistat filtrable, CRUD, filtres per estat/categoria"],
        ["Sol·licituds", "requests_page.dart", "Llistat per rol, accions acceptar/rebutjar/cancelar"],
        ["Detall sol·licitud", "request_detail_page.dart", "Vista completa amb materials i preu total"],
        ["Reserves", "reservations_page.dart", "Llistat amb accions aprovar/rebutjar/devolver/cancelar"],
        ["Detall reserva", "reservation_detail_page.dart", "Línies de material, danys, preu"],
        ["Calendari", "calendar_page.dart", "Vista mensual amb badges R/S, events del dia"],
        ["Usuaris", "users_page.dart", "Llistat filtrable per rol/actiu, CRUD"],
        ["Form Excursió", "excursion_form_page.dart", "Crear/editar amb tots els camps + categories múltiples"],
        ["Form Material", "equipment_form_page.dart", "Crear/editar amb stock, preu, dany, categories, estat"],
        ["Form Sol·licitud", "request_form_page.dart", "Selecció excursió, participants, recàlcul materials automàtic"],
        ["Form Reserva", "reservation_form_page.dart", "Dates, línies (equip + quantitat), danys"],
        ["Form Usuari", "user_form_page.dart", "Dades personals, rol, estat"],
    ]
)

body("Gestió d'estat amb Riverpod:", bold=True)
body(
    "Cada recurs (excursions, equipament, sol·licituds, reserves, usuaris) té "
    "un AsyncNotifierProvider que simula les operacions CRUD amb un delay "
    "artificial (800ms càrrega, 500ms acció). A més, cada recurs té un provider "
    "de filtrat (.family) que combina múltiples filtres (query, estat, categoria, "
    "rang de dates, idUsuari) de forma declarativa."
)
body(
    "Els resolvers (resolvers_provider.dart) permeten obtenir noms i imatges "
    "d'entitats relacionades (ex: nom de l'excursió d'una sol·licitud) de forma "
    "eficient amb Provider.family."
)

body("Serveis de negoci:", bold=True)
bullet("PricingService: calcularPrecioSolicitud (excursió × participants + materials × dies) i calcularCargoDanios (càrrec × items danyats).")
bullet("Resolvers: funcions pures per resoldre noms i imatges creuades entre entitats.")

img_placeholder("CAPTURES: Home Client (mòbil) + Home Admin (mòbil) en tema clar", h=8)
img_placeholder("CAPTURES: Catàleg d'excursions amb bottom sheet de filtres obert", h=8)
img_placeholder("CAPTURES: Formulari de nova excursió (camps, xips de categoria, selectors de data/hora)", h=8)
img_placeholder("CAPTURES: Catàleg de material + Formulari de material", h=8)
img_placeholder("CAPTURES: Sol·licituds amb accions + Formulari de sol·licitud (amb recàlcul automàtic)", h=8)
img_placeholder("CAPTURES: Reserves amb línies + Diàleg de registrar devolució", h=8)
img_placeholder("CAPTURES: Calendari mensual amb badges i llista d'events del dia", h=8)
img_placeholder("CAPTURES: Gestió d'usuaris + Formulari d'usuari", h=8)

heading2("3.4 Disseny")
body(
    "El tema visual es defineix en tres fitxers centralitzats a lib/app/theme/:"
)
body("app_colors.dart:", bold=True)
body(
    "Classe estàtica AppColors amb 22 constants de color (11 per tema clar, "
    "11 per tema fosc). Cada parell de colors manté coherència visual i contrast."
)
body("app_text_styles.dart:", bold=True)
body(
    "Classe estàtica amb 10 estils tipogràfics: titleLarge (26pt, w800, spacing 4), "
    "headlineSmall (20pt, bold), titleMedium (16pt, w600), bodyLarge (16pt, w700), "
    "bodyMedium (14pt), bodySmall (12pt), labelLarge (15pt, w600, spacing 1.1), "
    "labelMedium (13pt), labelSmall (11pt), titleSmall (9pt, w500)."
)
body("app_theme.dart:", bold=True)
body(
    "Dos ThemeData complets (light i dark) que configuren: scaffoldBackgroundColor, "
    "AppBarTheme (fons surfaceContainer, text onPrimary), ColorScheme complet "
    "amb els 11 colors, i TextTheme amb els 10 estils."
)
body("Widgets reutilitzables (core/widgets/):", bold=True)
bullet("AddFab: botó flotant personalitzat amb borde.")
bullet("PrimaryButton, SecondaryButton, TertiaryButton, MiniButton: 4 variants de botó.")
bullet("AppChoiceChip, AppFilterChip, AppChipWrap: xips de selecció única i múltiple.")
bullet("AppDateSelector, AppTimeSelector: selectors visuals de data i hora.")
bullet("AppDropdownField: dropdown genèric tipat amb validació.")
bullet("AppInputField: TextFormField estilitzat amb underline.")
bullet("AppImagePickerField: previsualitzador + botó d'imatge.")
bullet("AppTag: etiqueta colorada per estats/categories.")
bullet("ConfirmDialog: diàleg de confirmació reutilitzable.")
bullet("DetailSection + DetailRow: seccions de detall amb icona + label + valor.")
bullet("FilterBottomSheet: panel de filtres genèric amb xips + dates.")

img_placeholder("CAPTURES: comparativa tema clar vs tema fosc en 3 pantalles (login, home, catàleg)", h=8)
img_placeholder("CAPTURA: catàleg de components (design system) mostrant botons, inputs i cards", h=7)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  4. MILLORES FUTURES
# ══════════════════════════════════════════════════════════════════════════════
heading1("4. Millores futures")

body("Backend — Prioritat alta:", bold=True)
bullet("Aplicar AuthGuard JWT a tots els endpoints protegits (ara estan oberts).")
bullet("Hashear la password al crear usuaris (UserService.create).")
bullet("Moure el secret JWT a variable d'entorn (.env).")
bullet("Crear mòduls Solicitud i Reserva al backend (ara només existeixen al frontend).")
bullet("Implementar refresh token i invalidació de sessions.")
bullet("Hash + salt de passwords amb cost factor configurable.")

body("Backend — Prioritat mitjana:", bold=True)
bullet("Endpoint POST /auth/register per a registre públic.")
bullet("Forgot/Reset password per email (amb integració SMTP).")
bullet("Pujada d'imatges (integració amb S3, Cloudinary o similar).")
bullet("Paginació, sorting i filtres avançats a tots els endpoints GET.")
bullet("Role-based access control (RBAC) amb decorator personalitzat.")
bullet("Logging estructurat (Winston/Pino).")
bullet("Rate limiting amb @nestjs/throttler.")

body("Frontend — Prioritat alta:", bold=True)
bullet("Substituir totes les dades fake per crides reals a la API via Dio.")
bullet("Activar AuthStorage (flutter_secure_storage) per a la persistència del token.")
bullet("Configurar DioClient amb interceptor que afegeixi el Bearer token automàticament.")
bullet("Gestió d'errors de xarxa amb missatges amigables.")

body("Frontend — Prioritat mitjana:", bold=True)
bullet("Pantalla de registre d'usuari.")
bullet("Recuperació de contrasenya.")
bullet("Notificacions push (Firebase Cloud Messaging).")
bullet("Mapa integrat (Google Maps) per a punts d'inici/fi.")
bullet("Valoracions i comentaris d'excursions.")
bullet("Exportació a PDF de reserves i factures.")
bullet("Animacions i transicions entre pantalles.")
bullet("Tests unitaris i de widget.")

body("Negoci i producte:", bold=True)
bullet("Integració amb passarel·la de pagament (Stripe).")
bullet("Dashboard d'estadístiques per a admin (gràfics, KPIs).")
bullet("Versió web per a gestió d'escriptori.")
bullet("Multi-tenant (una instància per empresa client → SaaS).")
bullet("Internacionalització completa (ca, es, en, fr).")
bullet("App de guia (versió lite per a experts en ruta).")

page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  5. CONCLUSIONS
# ══════════════════════════════════════════════════════════════════════════════
heading1("5. Conclusions")

heading2("5.1 Personals")
body(
    "Ha estat el primer cop que he desenvolupat una aplicació completa full-stack "
    "des de zero, amb un frontend multiplataforma i un backend REST amb base de dades "
    "relacional. El repte més gran ha sigut gestionar el temps: mantenir el ritme "
    "entre frontend i backend alhora, decidir què prioritzar i què deixar per després."
)
body(
    "Flutter m'ha sorprès per bé. El sistema de widgets és molt potent i Riverpod, "
    "tot i que té una corba d'aprenentatge notable, un cop l'entens fa que la gestió "
    "d'estat sigui molt neta. He après a pensar en termes de providers i notifiers, "
    "i a separar la lògica de la UI de manera real, no només teòrica."
)
body(
    "De NestJS m'emporto la importància de la modularitat. Tenir cada recurs aïllat "
    "en el seu mòdul (controller + service + DTO) fa que el codi sigui fàcil d'entendre "
    "i de mantenir. Prisma és molt còmode per treballar amb la BD sense escriure SQL, "
    "i Swagger et documenta la API gairebé sol."
)
body(
    "El que canviaria si tornés a començar: hauria connectat el frontend al backend "
    "molt abans. Treballar amb dades fake durant tant de temps va fer que la integració "
    "final fos una fase a part, quan hauria d'haver estat contínua."
)

heading2("5.2 Tècniques")
body(
    "Tècnicament, el projecte ha assolit els objectius principals: una API REST "
    "funcional amb autenticació, una app Flutter amb 17 pantalles completes, "
    "un sistema de tema dual, un design system intern i un entorn de dev containeritzat."
)
body("Llistat d'assoliments tècnics:")
bullet("Clean Architecture al frontend amb separació de capes (domain, data, presentation, services).")
bullet("Riverpod amb AsyncNotifier per a CRUD reactiu i filtrat combinat amb Provider.family.")
bullet("NestJS modular amb 7 mòduls REST, validació amb DTOs i documentació Swagger completa.")
bullet("Prisma ORM amb relacions M:N, seeds i migracions.")
bullet("Docker Compose per a un entorn de dev reproducible en un sol comando.")
bullet("Design system intern (catalog/) com a storybook per validar components.")
bullet("4 variants de login per explorar opcions de UI.")
bullet("Sistema de preus (PricingService) amb càlcul automàtic de sol·licituds i danys.")
body(
    "El punt feble tècnic és la falta de connexió real entre front i back (el frontend "
    "funciona amb fakes) i la manca de guards d'autenticació al backend. Però la base "
    "arquitectònica és sòlida per afegir-ho sense grans refactors."
)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  6. REFERÈNCIES BIBLIOGRÀFIQUES
# ══════════════════════════════════════════════════════════════════════════════
heading1("6. Referències bibliogràfiques")

refs = [
    "Flutter Documentation. (2024). Build apps for any screen. https://flutter.dev/docs",
    "Dart Documentation. (2024). Dart programming language. https://dart.dev/docs",
    "Riverpod. (2024). A reactive state-management library for Dart/Flutter. https://riverpod.dev",
    "NestJS Documentation. (2024). A progressive Node.js framework. https://docs.nestjs.com",
    "Prisma Documentation. (2024). Next-generation ORM for Node.js & TypeScript. https://www.prisma.io/docs",
    "PostgreSQL. (2024). The world's most advanced open source relational database. https://www.postgresql.org/docs/16/",
    "Material Design 3. (2024). Design system by Google. https://m3.material.io",
    "Docker Documentation. (2024). Containerization platform. https://docs.docker.com",
    "JWT.io. (2024). Introduction to JSON Web Tokens. https://jwt.io/introduction",
    "OpenAPI Specification (Swagger). (2024). API description format. https://swagger.io/specification/",
    "bcrypt. (2024). Password hashing for Node.js. https://www.npmjs.com/package/bcrypt",
    "class-validator. (2024). Decorative validation for TypeScript. https://github.com/typestack/class-validator",
    "Dio. (2024). Powerful HTTP client for Dart. https://pub.dev/packages/dio",
    "table_calendar. (2024). Calendar widget for Flutter. https://pub.dev/packages/table_calendar",
    "flutter_secure_storage. (2024). Secure storage plugin. https://pub.dev/packages/flutter_secure_storage",
]
for i, ref in enumerate(refs, 1):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(f"[{i}] {ref}")
    run.font.name = FONT_BODY
    run.font.size = Pt(10)


# ── Save ───────────────────────────────────────────────────────────────────────
output = "memoria_outventura_v2.docx"
doc.save(output)
print(f"Document generat: {output}")
